"""
Document parsing service.
Handles PDF (text + tables + OCR), TXT, and DOCX files.
"""
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_number: int
    text: str
    tables: list = field(default_factory=list)
    is_ocr: bool = False


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    pages: list  # List[ParsedPage]
    page_count: int = 0

    def __post_init__(self):
        self.page_count = len(self.pages)

    @property
    def full_text(self) -> str:
        parts = []
        for page in self.pages:
            parts.append(page.text)
            for table in page.tables:
                parts.append(_table_to_markdown(table))
        return "\n\n".join(p for p in parts if p.strip())


def _table_to_markdown(table: list) -> str:
    """Convert a list-of-lists table to markdown format."""
    if not table or not table[0]:
        return ""

    rows = []
    header = table[0]
    rows.append("| " + " | ".join(str(c or "") for c in header) + " |")
    rows.append("| " + " | ".join("---" for _ in header) + " |")

    for row in table[1:]:
        cells = [str(c or "") for c in row]
        while len(cells) < len(header):
            cells.append("")
        rows.append("| " + " | ".join(cells[:len(header)]) + " |")

    return "\n".join(rows)


def parse_file(file_path: str) -> ParsedDocument:
    """Parse a file into structured pages with text and tables."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".txt":
        return _parse_txt(path)
    elif ext == ".docx":
        return _parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _parse_pdf(path: Path) -> ParsedDocument:
    """Parse PDF with text extraction, table detection, and OCR fallback."""
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []

            is_ocr = False
            # OCR fallback: if text is too sparse relative to page area, likely scanned
            if len(text.strip()) < 50 and page.width and page.height:
                ocr_text = _ocr_page(page)
                if ocr_text:
                    text = ocr_text
                    is_ocr = True

            pages.append(ParsedPage(
                page_number=i + 1,
                text=text,
                tables=tables,
                is_ocr=is_ocr,
            ))

    return ParsedDocument(filename=path.name, file_type="pdf", pages=pages)


def _ocr_page(page) -> str:
    """OCR a pdfplumber page using Tesseract."""
    try:
        import pytesseract
        from PIL import Image

        img = page.to_image(resolution=300)
        pil_image = img.original
        text = pytesseract.image_to_string(pil_image)
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR failed for page: {e}")
        return ""


def _parse_txt(path: Path) -> ParsedDocument:
    """Parse plain text file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(
        filename=path.name,
        file_type="txt",
        pages=[ParsedPage(page_number=1, text=text)],
    )


def _parse_docx(path: Path) -> ParsedDocument:
    """Parse DOCX file with text and table extraction."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))

    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            tables.append(rows)

    return ParsedDocument(
        filename=path.name,
        file_type="docx",
        pages=[ParsedPage(
            page_number=1,
            text="\n".join(full_text),
            tables=tables,
        )],
    )
